# Imports
import torch
import torch.nn as nn
import torch.optim as optim
import torch.nn.functional as F
import torchvision
from torchvision import datasets, transforms
from torch.utils.tensorboard import SummaryWriter
from torch.utils.data import DataLoader
from transformers import ViTForImageClassification

from sklearn.metrics import confusion_matrix, ConfusionMatrixDisplay
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

import time
# https://colab.research.google.com/github/NielsRogge/Transformers-Tutorials/blob/master/VisionTransformer/Fine_tuning_the_Vision_Transformer_on_CIFAR_10_with_the_%F0%9F%A4%97_Trainer.ipynb#scrollTo=Ik1n1OArMTk4
# import for docx
from docx import Document, table

# VGG16 Architecture
VGG11 = [64, 'M', 128, 'M', 256, 256, 'M', 512, 512, 'M', 512, 512, 'M']
#
VGG13 = [64,64,'M',128,128,'M',256,256,'M',512,512,'M',512,512,'M']
#
VGG16 = [64,64,'M',128,128,'M',256,256,256,'M',512,512,512,'M',512,512,512,'M']
#
VGG19 = [64,64,'M',128,128,'M',256,256,256,256,'M',512,512,512,512,'M',512,512,512,512,'M']
#

class VGG_Net(nn.Module):
    def __init__(self, in_channels=3, num_classes=10):
        super(VGG_Net, self).__init__()
        self.in_channels = in_channels
        #self.conv_layers = self.create_conv_layers(VGG11)
        #self.conv_layers = self.create_conv_layers(VGG13)
        self.conv_layers = self.create_conv_layers(VGG16)
        #self.conv_layers = self.create_conv_layers(VGG19)

        self.fcs = nn.Sequential(
            nn.Linear(512*1*1, 32*32*3), # data is 32 x32 pixel then, 32 -> 5 x Maxpool Layer = 32->16->8->4->2->1
            nn.ReLU(),
            nn.Dropout(p=0.5),
            nn.Linear(32*32*3, 1536),
            nn.ReLU(),
            nn.Dropout(p=0.5),
            nn.Linear(1536, num_classes)
        )
        #self._initialize_weights()
    def forward(self, x):
        x = self.conv_layers(x)
        x = x.reshape(x.shape[0], -1)
        x = self.fcs(x)
        return x

    def create_conv_layers(self, architecture):
        layers = []
        in_channels = self.in_channels

        for x in architecture:
          if type(x) == int:
            outchannels = x
            layers += [nn.Conv2d(in_channels=in_channels, out_channels=outchannels, kernel_size=3,
                                 stride=1, padding=1),
                       nn.BatchNorm2d(x),
                       nn.ReLU()]
            in_channels = x
          elif x == 'M':
            layers += [nn.MaxPool2d(kernel_size=2, stride=2)]

        return nn.Sequential(*layers)

    def _initialize_weights(self):
      for m in self.modules():
        if isinstance(m, nn.Linear):
          nn.init.kaiming_normal_(m.weight, mode='fan_out', nonlinearity='relu')
          if m.bias is not None:
            nn.init.constant_(m.bias, 0)

def docwriter(txtlst):
    Report_Doc = Document()
    Report_Doc.add_heading('Begin Report', 0)
    for lin in txtlst:
        Report_Doc.add_paragraph(lin, style='List Number')
    Report_Doc.save('VGG16 Result Summary.docx')



#standardisation
standardisation = transforms.Compose([


    #ref: https://www.kaggle.com/code/vikasbhadoria/cifar10-high-accuracy-model-build-on-pytorch
    #transforms.Resize((32,32)),  #resises the image so it can be perfect for our model.
    transforms.RandomHorizontalFlip(), # FLips the image w.r.t horizontal axis
    #transforms.RandomRotation(10),     #Rotates the image to a specified angel
    transforms.RandomAffine(0, shear=10, scale=(0.8,1.2)), #Performs actions like zooms, change shear angles.
    #transforms.ColorJitter(brightness=0.2, contrast=0.2, saturation=0.2), # Set the color params

    transforms.ToTensor(),
    #transforms.Normalize(mean=[0.4914, 0.4822, 0.4465], std = [0.2023, 0.1994, 0.2010])
    transforms.Normalize(mean=[0.4913, 0.4821, 0.4465], std = [0.2470, 0.2434, 0.2615])
])

# load data
train_dataset = datasets.CIFAR10(root='./data', train=True, download=True, transform=standardisation)
test_dataset = datasets.CIFAR10(root='./data', train=False, download=True, transform=standardisation)

BATCH_SIZE = 64
train_loader = DataLoader(dataset = train_dataset, batch_size = BATCH_SIZE, shuffle=True)
test_loader = DataLoader(dataset = test_dataset, batch_size = BATCH_SIZE, shuffle=True)
#print('train:', train_loader.dataset.data.shape) # 32 by 32 pictures with RGB channels
#print('test:', test_loader.dataset.data.shape)

if torch.cuda.is_available():
    device = torch.device('cuda')
else:

#  model = VGG_Net(in_channels = 3, num_classes=10).to(device) # for training model using no function
model = ViTForImageClassification.from_pretrained('google/vit-base-patch16-224-in21k',
                                                  id2label=id2label,
                                                  label2id=label2id)
loss_fn = nn.CrossEntropyLoss()
optimizer = optim.Adam(model.parameters(), lr=0.0001)
#optimizer = torch.optim.SGD(model.parameters(), lr = 0.0001, momentum=0.9,weight_decay=5e-4)
lr_scheduler = torch.optim.lr_scheduler.StepLR(optimizer, step_size=10, gamma=0.1)



#for i in model.named_children():
#    print(i)

total_batch = len(train_loader)
total_loss =0
loss_arr = []
txt = []
epochs = 3

start_time = time.time()
txt.append(start_time)
for epoch in range(epochs):
    epoch_start = time.time()
    avg_loss = 0
    # lr_scheduler.step() # Moved this down
    '''we are getting the warning: /usr/local/lib/python3.10/dist-packages/torch/optim/lr_scheduler.py:143:
    UserWarning: Detected call of `lr_scheduler.step()` before `optimizer.step()`.
    In PyTorch 1.1.0 and later, you should call them in the opposite order: `optimizer.step()` before `lr_scheduler.step()`.
    Failure to do this will result in PyTorch skipping the first value of the learning rate schedule.
    See more details at https://pytorch.org/docs/stable/optim.html#how-to-adjust-learning-rate
    warnings.warn("Detected call of `lr_scheduler.step()` before `optimizer.step()`. " '''

    for num, (image, label) in enumerate(train_loader):
        x = image.to(device)
        y_ = label.to(device)

        optimizer.zero_grad()
        output = model.forward(x)

        loss = loss_fn(output, y_)
        loss.backward()
        optimizer.step()

        avg_loss += loss / total_batch

    loss_arr.append(avg_loss)
    lr_scheduler.step()
    txt.append(f'[Epoch: {epoch+1} / {epochs}, loss = {avg_loss}], epoch time: {round(time.time()-epoch_start, 2)} seconds')
    #print(f'[Epoch: {epoch+1} / {epochs}, loss = {avg_loss}], epoch time: {round(time.time()-epoch_start, 2)} seconds')


txt.append(f'Training time: {round(time.time()-start_time,2)})')
txt.append('Training Finished !')

docwriter(txt)
#print(f'Training time: {round(time.time()-start_time,2)})')
#print('Training Finished !')

correct = 0
total = 0
predictions=[]
actual_labels=[]

with torch.no_grad():
    for num, (image, label) in enumerate(test_loader):
        x = image.to(device)
        y_ = label.to(device)
        outputs = model.forward(x)

        _, predicted = torch.max(outputs.data, 1)

        total += y_.size(0)
        correct += (predicted == y_).sum().item()
        predictions.extend(predicted.cpu().numpy())
        actual_labels.extend(y_.cpu().numpy())

results_df=pd.DataFrame({"Actual":actual_labels,"Predicted":predictions})
#print(results_df)

#print(f'Accuracy: {correct / total * 100}%')

# reference: https://www.kaggle.com/code/muhammedtuncayaydn/96-accuracy-with-pytorch-vgg16-model
from sklearn.metrics import accuracy_score

true_labels=results_df["Actual"]
predicted_labels=results_df["Predicted"]

accuracy=accuracy_score(true_labels,predicted_labels)

correct_predictions = (true_labels == predicted_labels).sum()