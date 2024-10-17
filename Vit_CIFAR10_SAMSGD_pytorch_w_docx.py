# Summary: We coded ViT from scratch to implement it on CIFAR10.
# Maximum accuracy was __ after __ Epoch and time was __.

from pathlib import Path
import torch
#import pytorch_lightning as pl

import numpy as np
import math
from tqdm import tqdm, trange
from torch.utils.tensorboard import SummaryWriter


import torch.nn as nn
from torch.optim import Adam, AdamW
from torch.nn import CrossEntropyLoss
from torch.utils.data import Dataset, DataLoader

from torchvision import datasets
from torchvision.datasets import CIFAR10
from torchvision.transforms import transforms, RandAugment #,AutoAugment, AutoAugmentPolicy, ToTensor



from sklearn.metrics import confusion_matrix, ConfusionMatrixDisplay
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

import datetime


# import for docx
from docx import Document, table

#standardisation
standardisation = transforms.Compose([
    transforms.RandomHorizontalFlip(),
    #transforms.AutoAugment(),
    #RandAugment(),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.4913, 0.4821, 0.4465], std = [0.2470, 0.2434, 0.2615])
])

# load data
train_dataset = datasets.CIFAR10(root='./data', train=True, download=True, transform=standardisation)
test_dataset = datasets.CIFAR10(root='./data', train=False, download=True, transform=standardisation)

# hyperparameters
batch_size = 64
N_EPOCHS = 30
LR = 0.0001

#data
train_loader = DataLoader(dataset = train_dataset, batch_size = batch_size , shuffle=True)
test_loader = DataLoader(dataset = test_dataset, batch_size = batch_size , shuffle=True)

#device
if torch.cuda.is_available():
    device = torch.device('cuda')
else:
    device = torch.device('cpu')

# Create patches

def patchify(images, n_patches):
  n, c, h, w = images.shape

  assert h == w, "Patchify method is implemented for square images only"

  patches = torch.zeros(n, n_patches ** 2, h * w * c // n_patches ** 2)
  patch_size = h // n_patches

  for idx, image in enumerate(images):
      for i in range(n_patches):
          for j in range(n_patches):
              patch = image[:, i * patch_size: (i + 1) * patch_size, j * patch_size: (j + 1) * patch_size]
              patches[idx, i * n_patches + j] = patch.flatten()
  return patches


# Positional encoding
def get_positional_encoding(n_patches, hidden_d):
  result = torch.ones(n_patches, hidden_d)
  for i in range(n_patches):
    for j in range(hidden_d):
      result[i][j] = math.sin(i / (10000 ** (j / hidden_d))) if j % 2 == 0 else math.cos(i / (10000 ** ((j - 1) / hidden_d)))
  return result

#input embedding
class MyVit(nn.Module):
  def __init__(self, chw, n_patches=8, n_blocks=2, hidden_d=8, n_heads=2, out_d=10):
    super(MyVit, self).__init__() # this will give access to the methods and properties of parents and sibling class

    #Attributes
    self.chw = chw # ( C , H , W )
    self.n_patches = n_patches
    self.n_blocks = n_blocks
    self.n_heads = n_heads
    self.hidden_d = hidden_d

    assert chw[1] % n_patches == 0 and chw[2] % n_patches == 0, "Patchify method is implemented for square images only"
    self.patch_size = (chw[1] / n_patches, chw[2] / n_patches)

    # Linear Mapper
    self.input_d = int(chw[0]*self.patch_size[0]*self.patch_size[1])
    self.linear_mapper = nn.Linear(self.input_d, self.hidden_d)

    # Token Classification
    self.class_token = nn.Parameter(torch.randn(1, 1, self.hidden_d))

   # 3) Positional embedding
    self.register_buffer(
        "positional_embeddings",
        get_positional_encoding(self.n_patches**2 + 1, self.hidden_d),
        persistent=False,
    )

    # 4) Transformer encoder blocks
    self.blocks = nn.ModuleList(
        [MyViTBlock(self.hidden_d, self.n_heads) for _ in range(self.n_blocks)]
    )

    # 5) Classification MLPk
    self.mlp = nn.Sequential(nn.Linear(self.hidden_d, out_d), nn.Softmax(dim=-1))

  def forward(self, images):
    # Dividing images into patches
    n, c, h, w = images.shape
    patches = patchify(images, self.n_patches).to(self.positional_embeddings.device)

    # Running linear layer tokenization
    # Map the vector corresponding to each patch to the hidden size dimension
    tokens = self.linear_mapper(patches)

    # Adding classification token to the tokens
    tokens = torch.cat((self.class_token.expand(n, 1, -1), tokens), dim=1)

    # Adding positional embedding
    out = tokens + self.positional_embeddings.repeat(n, 1, 1)

    # Transformer Blocks
    for block in self.blocks:
        out = block(out)

    # Getting the classification token only
    out = out[:, 0]

    return self.mlp(out) # Map to output dimension, output category distribution

    return out
  
#Multihead attention block
class MyMSA(nn.Module):
    def __init__(self, d, n_heads=2):
        super(MyMSA, self).__init__()
        self.d = d
        self.n_heads = n_heads

        assert d % n_heads == 0, f"Can't divide dimension {d} into {n_heads} heads"

        d_head = int(d / n_heads)
        self.q_mappings = nn.ModuleList([nn.Linear(d_head, d_head) for _ in range(self.n_heads)])
        self.k_mappings = nn.ModuleList([nn.Linear(d_head, d_head) for _ in range(self.n_heads)])
        self.v_mappings = nn.ModuleList([nn.Linear(d_head, d_head) for _ in range(self.n_heads)])
        self.d_head = d_head
        self.softmax = nn.Softmax(dim=-1)

    def forward(self, sequences):
        # Sequences has shape (N, seq_length, token_dim)
        # We go into shape    (N, seq_length, n_heads, token_dim / n_heads)
        # And come back to    (N, seq_length, item_dim)  (through concatenation)
        result = []
        for sequence in sequences:
            seq_result = []
            for head in range(self.n_heads):
                q_mapping = self.q_mappings[head]
                k_mapping = self.k_mappings[head]
                v_mapping = self.v_mappings[head]

                seq = sequence[:, head * self.d_head: (head + 1) * self.d_head]
                q, k, v = q_mapping(seq), k_mapping(seq), v_mapping(seq)

                attention = self.softmax(q @ k.T / (self.d_head ** 0.5))
                seq_result.append(attention @ v)
            result.append(torch.hstack(seq_result))
        return torch.cat([torch.unsqueeze(r, dim=0) for r in result])
    

#residual block
class MyViTBlock(nn.Module):
    def __init__(self, d, n_heads):
        super(MyViTBlock, self).__init__()
        self.hidden_d = d
        self.n_heads = n_heads

        self.norm1 = nn.LayerNorm(d)
        self.mhsa = MyMSA(d, n_heads)
        self.norm2 = nn.LayerNorm(d)
        self.mlp = nn.Sequential(
            nn.Linear(d, 4 * d),
            nn.GELU(),
            nn.Linear(4 * d, d)
        )

    def forward(self, x):
        x = x + self.mhsa(self.norm1(x))
        x = x + self.mlp(self.norm2(x))
        return x
    

# SAM block  
# From Davda54
import torch


class SAM(torch.optim.Optimizer):
    def __init__(self, params, base_optimizer, rho=0.05, adaptive=False, **kwargs):
        assert rho >= 0.0, f"Invalid rho, should be non-negative: {rho}"

        defaults = dict(rho=rho, adaptive=adaptive, **kwargs)
        super(SAM, self).__init__(params, defaults)

        self.base_optimizer = base_optimizer(self.param_groups, **kwargs)
        self.param_groups = self.base_optimizer.param_groups
        self.defaults.update(self.base_optimizer.defaults)

    @torch.no_grad()
    def first_step(self, zero_grad=False):
        grad_norm = self._grad_norm()
        for group in self.param_groups:
            scale = group["rho"] / (grad_norm + 1e-12)

            for p in group["params"]:
                if p.grad is None: continue
                self.state[p]["old_p"] = p.data.clone()
                e_w = (torch.pow(p, 2) if group["adaptive"] else 1.0) * p.grad * scale.to(p)
                p.add_(e_w)  # climb to the local maximum "w + e(w)"

        if zero_grad: self.zero_grad()

    @torch.no_grad()
    def second_step(self, zero_grad=False):
        for group in self.param_groups:
            for p in group["params"]:
                if p.grad is None: continue
                p.data = self.state[p]["old_p"]  # get back to "w" from "w + e(w)"

        self.base_optimizer.step()  # do the actual "sharpness-aware" update

        if zero_grad: self.zero_grad()

    @torch.no_grad()
    def step(self, closure=None):
        assert closure is not None, "Sharpness Aware Minimization requires closure, but it was not provided"
        closure = torch.enable_grad()(closure)  # the closure should do a full forward-backward pass

        self.first_step(zero_grad=True)
        closure()
        self.second_step()

    def _grad_norm(self):
        shared_device = self.param_groups[0]["params"][0].device  # put everything on the same device, in case of model parallelism
        norm = torch.norm(
                    torch.stack([
                        ((torch.abs(p) if group["adaptive"] else 1.0) * p.grad).norm(p=2).to(shared_device)
                        for group in self.param_groups for p in group["params"]
                        if p.grad is not None
                    ]),
                    p=2
               )
        return norm

    def load_state_dict(self, state_dict):
        super().load_state_dict(state_dict)
        self.base_optimizer.param_groups = self.param_groups


def docwriter(txtlst,img):
    Report_Doc = Document()
    Report_Doc.add_heading('Begin Report', 0)
    for lin in txtlst:
        Report_Doc.add_paragraph(lin, style='List Number')
    Report_Doc.add_picture(img)
    Report_Doc.save('ViT Result Summary with SAM_SGD.docx')

#Define model    
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model = MyVit((3, 32, 32), n_patches=2, n_blocks=12, hidden_d=480, n_heads=12, out_d=10).to(device)

txt =[]
#get model parameters
pytorch_total_params = sum(p.numel() for p in model.parameters() if p.requires_grad)

#print(f"Total params: {pytorch_total_params}")
txt.append(f"Total params: {pytorch_total_params}")

#Final Run
Begin_Time = datetime.datetime.now()

#get model parameters
pytorch_total_params = sum(p.numel() for p in model.parameters() if p.requires_grad)

#print(f"Total params: {pytorch_total_params}")
txt.append(f"Total params: {pytorch_total_params}")

# Training loop
optimizer = Adam(model.parameters(), lr=LR)
base_optimizer = torch.optim.SGD  # define an optimizer for the "sharpness-aware" update
optimizer = SAM(model.parameters(), base_optimizer, lr=0.1, momentum=0.9)
lr_scheduler = torch.optim.lr_scheduler.StepLR(optimizer, step_size=10, gamma=0.1)
#loss_fn = CrossEntropyLoss()
#lr_scheduler = torch.optim.lr_scheduler.StepLR(optimizer, step_size=10, gamma=0.1)
criterion = CrossEntropyLoss()


for epoch in trange(N_EPOCHS, desc="Training"):
    train_loss = 0.0
    #lr_scheduler.step()    
    #for batch in tqdm(train_loader, desc=f"Epoch {epoch + 1} in training", leave=False):
    for num, (image, label) in enumerate(train_loader):
        x = image.to(device)
        y = label.to(device)    

        inputs = image.to(device)
        targets = label.to(device)


        # first forward-backward step
        predictions = model(inputs)
        loss = criterion(predictions, targets)
        loss.mean().backward()
        optimizer.first_step(zero_grad=True)

        # second forward-backward step
        criterion(model(inputs), targets).mean().backward()
        optimizer.second_step(zero_grad=True)

        train_loss += loss.item() / len(train_loader)
    
    
    
    txt.append(f"Epoch {epoch + 1}/{N_EPOCHS} loss: {train_loss:.2f}")
    lr_scheduler.step()
txt.append(f"Training time: {datetime.datetime.now() - Begin_Time}")



## Test loop

predictions=[]
actual_labels=[]
with torch.no_grad():
    correct, total = 0, 0
    test_loss = 0.0
    for batch in tqdm(test_loader, desc="Testing"):
        x, y = batch
        x, y = x.to(device), y.to(device)
        y_hat = model(x)
        loss = criterion(y_hat, y)
        test_loss += loss.detach().cpu().item() / len(test_loader)

        correct += torch.sum(torch.argmax(y_hat, dim=1) == y).detach().cpu().item()
        total += len(x)

        _, predicted = torch.max(y_hat.data, 1)
        #total += y.size(0)
        predictions.extend(predicted.cpu().numpy())
        actual_labels.extend(y.cpu().numpy())

    
    results_df=pd.DataFrame({"Actual":actual_labels,"Predicted":predictions})
    
txt.append(f"Test loss: {test_loss:0.2f}")
txt.append(f"Test accuracy: {correct / total * 100:.2f}%")
import seaborn as sns
from sklearn.metrics import confusion_matrix

true_labels=results_df["Actual"]
predicted_labels=results_df["Predicted"]

confusionMatrix = confusion_matrix(true_labels, predicted_labels)
 
x_axis_labels = ["airplanes", "cars", "birds", "cats", "deer", "dogs", "frogs", "horses", "ships","trucks"] # labels for x-axis
y_axis_labels = ["airplanes", "cars", "birds", "cats", "deer", "dogs", "frogs", "horses", "ships","trucks"] # labels for y-axis
cnf_htmap =sns.heatmap(confusionMatrix, xticklabels=x_axis_labels, yticklabels=y_axis_labels, annot=True, cmap='viridis', fmt='g')
plt.xlabel('Actual', fontsize = 15) # x-axis label with fontsize 15
plt.ylabel('Predicted', fontsize = 15) # y-axis label with fontsize 15
cnf_htmap.figure.savefig("ht_map.png")
docwriter(txt,"ht_map.png")